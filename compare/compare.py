# -*- coding: utf-8 -*-
"""
Created on Mon Oct 19 16:00:06 2020

@author: Nathanael Judge, FINTECH CONSULTANCY

license: Apache 2.0,

Note: only tested on windows 10
"""
import pandas as pd
import sys
import re
import os
import win32com.client
from docx import *

# Hardcoded for now. Wondering where configurations should go for future iterations
columns = ['Definition', 'Revised Definition', 'DeliveryID','Name', 'Category']

def checkExtension(filename:str, extension:str):
    string = ''
    for character in filename[-1::-1]:
        if character == '.':
            break
        string += character
    if ''.join(string[-1::-1]) != extension:
        return(False)
    return(True)

def generate_doc(filepath = 'Favorites.xlsx', output_name = 'Example.docx', ID = '1'):
    print(filepath)
    if type(filepath) != str:
        filepath = str(filepath)
    if type(output_name) != str:
        output_name = str(output_name)
    if type(ID) != str:
        ID = str(ID)
        
    if not checkExtension(filepath.split('/')[-1], 'xlsx'):
        print("1. Error -- Type of file accepted should be an excel file")
        return(False)
    else:
        print("1. Input file is of type xlsx.")
    if not checkExtension(output_name.split('/')[-1], 'docx'):
        print("2. Error -- Type of outputted should be a docx file")
        return(False)
    else:
        print("2. Output file is of type docx.")
    print("Opening the data")
    
    print("Warning: sheet name is 'Sheet1'")
    data = pd.read_excel(filepath, 'Sheet1')
    count, _ = data.shape
    print("Data has succesfully been opened")

    
    ''' Check will not work within python. Must be run from shell'''
    # if not os.path.exists('Deliverytemplate.docx'):
    #     print(f"3. Error -- Make sure Deliverytemplate.docx is in the same directory as the {os.path.basename(__file__)}")
    # else:
    #     print(f"3. Deliverytemplate.docx is in the same directory as the {os.path.basename(__file__)}")
    print("Writing Revised contnet")
    document = Document('Deliverytemplate.docx')
    document.add_heading('Manifest', level=1)
    p = document.add_paragraph('')
    p.add_run('Delivery Type: test \n').bold = True
    p.add_run(f'Number of fields: {count}').bold = True
    document.save("Revised.docx")

    print("Writing Original Content")
    original = Document('Deliverytemplate.docx')
    original .add_heading('Manifest', level=1)
    p = original.add_paragraph('')
    p.add_run('Delivery Type: test\n').bold = True
    p.add_run(f'Number of fields: {count}').bold = True
    original.save("Original.docx")
    
    p = document.add_paragraph('')
    po = original.add_paragraph('')
    for i in range(count):
        if pd.isna(data['Category'].iloc[i]):
            pass
        else:
            p.add_run("Category:").italic = True
            p.add_run(f" {data['Category'].iloc[i].strip()}\n")
            po.add_run("Category:").italic = True
            po.add_run(f" {data['Category'].iloc[i].strip()}\n")
        if pd.isna(data['Name'].iloc[i]):
            pass
        else:
            p.add_run('Name:').italic = True
            p.add_run(f" {data['Name'].iloc[i].strip()}\n")
            po.add_run("Name:").italic = True
            po.add_run(f" {data['Name'].iloc[i].strip()}\n")
        if pd.isna(data['Revised Definition'].iloc[i]):
            pass
        else:
            p.add_run(f" {data['Revised Definition'].iloc[i].strip()}\n")
            po.add_run(f" {data['Definition'].iloc[i].strip()}\n")
        p.add_run("\n")
        po.add_run("\n")
    document.save("Revised.docx")
    original.save("Original.docx")
            # conda install -c anaconda pywin32 fixed issue with this line
    Application=win32com.client.gencache.EnsureDispatch('Word.Application')
    Application.CompareDocuments(Application.Documents.Open(os.getcwd() + "\\Original.docx"),
                                 Application.Documents.Open(os.getcwd() + "\\Revised.docx"))
    Application.ActiveDocument.ActiveWindow.View.Type = 3
    print(os.getcwd() + f"\\{output_name}")
    Application.ActiveDocument.SaveAs (FileName = os.getcwd() + f"\\{output_name}")
    Application.Quit()
    return(True)

def main(argv):
    print("Generating Report file")
    file = re.sub('\\\\', '/', argv[0])
    file = fr'{file}'
    file2 = re.sub('\\\\', '/', argv[1])
    if generate_doc(file, file2, argv[2]):
        print("Report been generated")
    else:
        print("Report failed to generate")


if __name__ == "__main__":
    main(sys.argv[1:])
